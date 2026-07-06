"""Document writer — the assistant drafts, this module makes REAL files.

The local model produces plain text; these functions turn it into genuine
Office documents: .docx via python-docx and .xlsx via openpyxl, including
live SUM formulas that calculate when the file opens in Excel, Microsoft
365, or LibreOffice. Read/parse/write are separated so each piece is
unit-testable without the model or the GUI.

Demonstrable claim this module supports: "a fully local assistant that
drafts real Word and Excel documents on request."
"""
from __future__ import annotations

import os
import re
from datetime import datetime

try:
    import docx as _docx
except Exception:
    _docx = None
try:
    import openpyxl as _openpyxl
except Exception:
    _openpyxl = None

# System prompts that make a small local model behave like a document
# clerk. The explicit "this is your job" framing matters: bare requests
# ("write a letter to my landlord") can trip a 3B model's guardrails
# into refusing routine paperwork.
LETTER_SYSTEM = (
    "You are a document-drafting assistant inside the user's personal "
    "productivity app. Drafting the user's own routine personal and "
    "business correspondence and paperwork (letters, memos, summaries) "
    "is your job. Output only the requested document text, ready to "
    "paste — no commentary, no markdown."
)

SHEET_SYSTEM = (
    "You are a spreadsheet-drafting assistant inside the user's personal "
    "productivity app. You turn requests into simple tables. Output "
    "format, with NO other text:\n"
    "Line 1: a short title for the sheet\n"
    "Line 2: column headers separated by |\n"
    "Then one line per row, cells separated by |. Numbers must be plain "
    "(no $ signs, no thousands separators)."
)


def sheet_prompt(request: str) -> str:
    return ("Create this as a table: " + request.strip() +
            "\nRemember: title line, then header line, then data lines, "
            "cells separated by |. Numbers plain. No other text.")


def letter_prompt(request: str) -> str:
    return "Draft the following document for me: " + request.strip()


def looks_like_refusal(reply: str) -> bool:
    head = (reply or "").strip().lower()
    return head.startswith(("i cannot", "i can't", "i am not able",
                            "i'm not able", "i won't", "sorry, i"))


def parse_table(reply: str):
    """Model reply -> (title, headers, rows). Tolerant: skips malformed
    lines; numeric-looking cells become floats. Returns (None, [], [])
    when nothing usable came back."""
    lines = [ln.strip() for ln in (reply or "").splitlines() if ln.strip()]
    # Drop markdown fences/separators the model sometimes adds anyway.
    lines = [ln for ln in lines
             if not re.fullmatch(r"[`|\-\s:]+", ln)]
    if not lines:
        return None, [], []
    title = None
    if "|" not in lines[0]:
        title = lines[0].strip().strip('"')
        lines = lines[1:]
    piped = [ln for ln in lines if "|" in ln]
    if not piped:
        return title, [], []
    headers = [c.strip() for c in piped[0].split("|") if c.strip()]
    rows = []
    for ln in piped[1:]:
        cells = [c.strip() for c in ln.split("|")]
        cells = [c for c in cells if c != ""]
        if not cells:
            continue
        out = []
        for c in cells:
            cleaned = c.replace("$", "").replace(",", "")
            try:
                out.append(float(cleaned))
            except ValueError:
                out.append(c)
        rows.append(out)
    return title, headers, rows


def write_table_xlsx(path: str, title: str, headers: list, rows: list) -> str:
    """Write a real .xlsx: bold headers, data rows, and a TOTAL row with a
    live =SUM() formula under every numeric column. Returns the path."""
    if _openpyxl is None:
        raise RuntimeError("openpyxl is not installed")
    wb = _openpyxl.Workbook()
    ws = wb.active
    # Excel forbids []:*?/\ in sheet-tab names (a title like
    # "Budget: July" would abort the whole write).
    safe = re.sub(r"[\[\]:*?/\\]", " ", title or "Sheet1")
    safe = re.sub(r"\s+", " ", safe).strip()[:31]
    ws.title = safe or "Sheet1"
    if headers:
        ws.append(headers)
        for cell in ws[1]:
            cell.font = _openpyxl.styles.Font(bold=True)
    for r in rows:
        ws.append(r)
    # TOTAL row: a SUM formula for each column that is numeric in EVERY
    # data row it appears in (a column of prices, minutes, quantities…).
    if rows:
        ncols = max(len(r) for r in rows)
        first = 2 if headers else 1
        last = first + len(rows) - 1
        total_row = []
        any_total = False
        for ci in range(ncols):
            vals = [r[ci] for r in rows if ci < len(r)]
            if vals and all(isinstance(v, float) for v in vals):
                col = _openpyxl.utils.get_column_letter(ci + 1)
                total_row.append(f"=SUM({col}{first}:{col}{last})")
                any_total = True
            else:
                total_row.append("TOTAL" if not any_total and ci == 0 else "")
        if any_total:
            ws.append(total_row)
            for cell in ws[ws.max_row]:
                cell.font = _openpyxl.styles.Font(bold=True)
    # Readable column widths.
    for ci in range(1, (max((len(r) for r in rows), default=len(headers) or 1)) + 1):
        col = _openpyxl.utils.get_column_letter(ci)
        width = 12
        for row in ws.iter_rows(min_col=ci, max_col=ci, values_only=True):
            v = row[0]
            if v is not None:
                width = max(width, min(40, len(str(v)) + 2))
        ws.column_dimensions[col].width = width
    os.makedirs(os.path.dirname(path), exist_ok=True)
    wb.save(path)
    return path


def write_letter_docx(path: str, title: str, body: str) -> str:
    """Write a real .docx: optional heading + one paragraph per blank-line
    block of the drafted text. Returns the path."""
    if _docx is None:
        raise RuntimeError("python-docx is not installed")
    doc = _docx.Document()
    if title:
        doc.add_heading(title, level=1)
    for block in re.split(r"\n\s*\n", (body or "").strip()):
        if block.strip():
            doc.add_paragraph(block.strip())
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


def suggest_filename(kind: str, title: str | None, when: datetime | None = None) -> str:
    """Safe, dated filename: 'Rent letter 2026-07-04 1432.docx'."""
    when = when or datetime.now()
    base = (title or ("Letter" if kind == "docx" else "Spreadsheet")).strip()
    base = re.sub(r'[<>:"/\\|?*\n\r\t]+', " ", base)
    base = re.sub(r"\s+", " ", base).strip()[:60] or "Document"
    return f"{base} {when.strftime('%Y-%m-%d %H%M')}.{kind}"
